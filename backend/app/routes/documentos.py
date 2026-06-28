from fastapi import APIRouter, HTTPException, Depends
from sqlalchemy.orm import Session
from docx import Document
from datetime import datetime
import os
import sys
import unicodedata
import re
import subprocess
import shutil
import glob
import traceback
from typing import Optional
import fitz  # PyMuPDF

from database import get_db, SessionLocal  # mantido
from app.models.usuario import Usuario
from app.models.advogado import Advogado
from app.models.extrato import Extrato
from app.utils.google_drive import upload_pdf_to_drive
from app.utils.zapsign import enviar_documentos_consolidados_para_assinatura, build_advogado_webhook_url
from app.core.time import now_sp  # ✅ fuso America/Sao_Paulo

router = APIRouter()

# ========================= STORAGE HELPERS =========================
def _resolve_storage_root() -> str:
    root = os.getenv("STORAGE_ROOT")
    if root:
        return os.path.abspath(root)
    # fallback: diretório storage dentro do backend
    return os.path.abspath(os.path.join(os.getcwd(), "storage"))


def _ensure_storage_dir(subdir: str, extrato_id: Optional[str]) -> Optional[str]:
    if not extrato_id:
        return None
    base = _resolve_storage_root()
    target = os.path.join(base, subdir, str(extrato_id))
    try:
        os.makedirs(target, exist_ok=True)
    except Exception as e:
        print(f"[Docs][WARN] Não foi possível criar diretório de storage {target}: {e}")
        return None
    return target


def _persist_pdf_in_storage(extrato_id: Optional[str], pdf_path: str, *, label: str) -> None:
    if not extrato_id or not pdf_path or not os.path.exists(pdf_path):
        return
    subdir = os.getenv("STORAGE_ASSINATURAS_SUBDIR", "Assinaturas")
    dest_dir = _ensure_storage_dir(subdir, extrato_id)
    if not dest_dir:
        return
    dest_path = os.path.join(dest_dir, os.path.basename(pdf_path))
    try:
        if os.path.abspath(pdf_path) != os.path.abspath(dest_path):
            shutil.copy2(pdf_path, dest_path)
    except Exception as e:
        print(f"[Docs][WARN] Falha ao copiar {label} para storage ({dest_path}): {e}")


# ========================= UTIL / SLUG =========================
def slugify(texto: str) -> str:
    texto = unicodedata.normalize('NFD', texto)
    texto = texto.encode('ascii', 'ignore').decode('utf-8')
    texto = re.sub(r'\s+', '_', texto)
    texto = re.sub(r'[^a-zA-Z0-9_]', '', texto)
    return texto.lower() or "arquivo"

# ========================= CONVERSÃO DOCX -> PDF =========================
def _try_docx2pdf(input_docx: str, output_pdf: str) -> bool:
    # docx2pdf depende do Word (macOS/Windows). Em Linux pulamos direto para LibreOffice.
    if sys.platform.startswith("linux"):
        return False
    try:
        import docx2pdf  # requer Word (macOS/Windows)
    except Exception:
        return False
    out_dir = os.path.dirname(output_pdf) or "."
    os.makedirs(out_dir, exist_ok=True)
    try:
        docx2pdf.convert(input_docx, out_dir)
        base = os.path.splitext(os.path.basename(input_docx))[0]
        candidate = os.path.join(out_dir, base + ".pdf")
        if os.path.exists(candidate):
            if os.path.abspath(candidate) != os.path.abspath(output_pdf):
                os.replace(candidate, output_pdf)
            return True
    except Exception:
        return False
    return False

def _find_soffice_binary() -> Optional[str]:
    """
    Busca o executável do LibreOffice em diferentes ambientes (Linux, Windows WSL, macOS).
    """
    candidates = [
        os.getenv("SOFFICE_PATH"),
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/usr/lib/libreoffice/program/soffice",
        "/usr/lib64/libreoffice/program/soffice",
        "/usr/local/bin/soffice",
    ]

    for path in candidates:
        if not path:
            continue
        expanded = os.path.expanduser(path)
        if os.path.isfile(expanded) and os.access(expanded, os.X_OK):
            return expanded

    # Instalações típicas no Windows/WSL (mantém a lista curta para evitar falsos positivos).
    extra_windows_candidates = [
        "/mnt/c/Program Files/LibreOffice/program/soffice.exe",
        "/mnt/c/Program Files (x86)/LibreOffice/program/soffice.exe",
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
    ]
    for path in extra_windows_candidates:
        if path and os.path.exists(path):
            return path

    return None

def _try_soffice(input_docx: str, output_pdf: str) -> bool:
    out_dir = os.path.dirname(output_pdf) or "."
    os.makedirs(out_dir, exist_ok=True)
    soffice_path = _find_soffice_binary()
    if not soffice_path:
        print("[Docs][WARN] LibreOffice (soffice) não encontrado.")
        print(f"[Docs][WARN] SOFFICE_PATH env = {os.getenv('SOFFICE_PATH')}")
        return False

    cmd = [
        soffice_path,
        "--headless",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        out_dir,
        input_docx,
    ]
    try:
        result = subprocess.run(cmd, check=False, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if result.returncode != 0:
            print(f"[Docs][WARN] soffice retornou código {result.returncode}: {result.stderr.decode(errors='ignore')}")
    except Exception as exc:
        print(f"[Docs][ERRO] Falha ao executar soffice: {exc}")
        return False

    base = os.path.splitext(os.path.basename(input_docx))[0]
    candidate = os.path.join(out_dir, base + ".pdf")
    if os.path.exists(candidate):
        if os.path.abspath(candidate) != os.path.abspath(output_pdf):
            os.replace(candidate, output_pdf)
        return True
    return False

def _pdf_tem_imagem(pdf_path: str) -> bool:
    try:
        with fitz.open(pdf_path) as doc:
            for p in doc:
                if p.get_images():
                    return True
    except Exception:
        pass
    return False

def _rasterizar_pdf(orig_pdf: str, out_pdf: str, scale: float = 2.0) -> None:
    """Gera um novo PDF a partir do original, rasterizando cada página (bake visual)."""
    with fitz.open(orig_pdf) as src:
        dst = fitz.open()
        for i, page in enumerate(src):
            mat = fitz.Matrix(scale, scale)  # ~ 144dpi se scale=2 em base 72dpi
            pix = page.get_pixmap(matrix=mat, alpha=False)
            rect = fitz.Rect(0, 0, pix.width, pix.height)
            new_page = dst.new_page(width=rect.width, height=rect.height)
            new_page.insert_image(rect, stream=pix.tobytes("png"))
        dst.save(out_pdf)
        dst.close()

def converter_docx_para_pdf(caminho_docx: str, caminho_pdf: str, *, rasterizar_se_sem_img: bool = True) -> None:
    print(f"[Docs] Convertendo DOCX -> PDF: {caminho_docx} -> {caminho_pdf}")
    ok = _try_docx2pdf(caminho_docx, caminho_pdf) or _try_soffice(caminho_docx, caminho_pdf)
    if not ok or not os.path.exists(caminho_pdf):
        print(f"[Docs][ERRO] Conversão falhou: {caminho_pdf}")
        raise RuntimeError(f"PDF não gerado: {caminho_pdf}")

    # Se não detectamos XObjects de imagem, pode ser: (a) assinatura em vetor/shape; (b) conversor removeu imagem; (c) imagem linkada.
    # Não bloqueamos. Se for crucial gerar PDF com imagens (para preservar visual em qualquer pipeline), rasterizamos como último recurso.
    if rasterizar_se_sem_img and not _pdf_tem_imagem(caminho_pdf):
        try:
            raster_out = caminho_pdf.replace(".pdf", "_raster.pdf")
            _rasterizar_pdf(caminho_pdf, raster_out, scale=2.5)  # ~180dpi
            if os.path.exists(raster_out) and _pdf_tem_imagem(raster_out):
                os.replace(raster_out, caminho_pdf)
        except Exception as e:
            # apenas registra; seguimos com o PDF original
            print(f"[aviso] Falha ao rasterizar PDF sem imagens: {e}")

# ========================= GERAÇÃO DOCX (PLACEHOLDERS) =========================
def preencher_documento(modelo_path: str, dados: dict, nome_saida: str) -> str:
    doc = Document(modelo_path)
    
    # ✨ COMPATIBILIDADE CPF/CNPJ
    # Se não tiver 'cpf', usar 'cpf_cnpj'
    if not dados.get('cpf') and dados.get('cpf_cnpj'):
        dados['cpf'] = dados['cpf_cnpj']
    
    # ⚠️ NÃO definir padrão CPF - deixar vazio para identificar problema
    if not dados.get('tipo_documento'):
        # Se não vier do frontend, tentar buscar do banco pelo extrato_id
        extrato_id = dados.get('extrato_id')
        if extrato_id:
            try:
                from database import SessionLocal
                from app.models.extrato import Extrato
                db = SessionLocal()
                extrato = db.query(Extrato).filter(Extrato.id == extrato_id).first()
                if extrato and extrato.tipo_documento:
                    dados['tipo_documento'] = extrato.tipo_documento
                    print(f"✅ tipo_documento buscado do banco: {extrato.tipo_documento}")
                else:
                    dados['tipo_documento'] = 'CPF'  # fallback apenas se não achar
                    print(f"⚠️ tipo_documento não encontrado, usando CPF como padrão")
                db.close()
            except Exception as e:
                print(f"❌ Erro ao buscar tipo_documento do banco: {e}")
                dados['tipo_documento'] = 'CPF'
        else:
            # Se não tem extrato_id, detectar pelo tamanho do CPF/CNPJ
            cpf_cnpj = dados.get('cpf_cnpj', '')
            cpf_limpo = cpf_cnpj.replace('.', '').replace('-', '').replace('/', '')
            if len(cpf_limpo) == 14:
                dados['tipo_documento'] = 'CNPJ'
                print(f"✅ tipo_documento detectado automaticamente: CNPJ (14 dígitos)")
            else:
                dados['tipo_documento'] = 'CPF'
                print(f"✅ tipo_documento detectado automaticamente: CPF (11 dígitos)")
    
    # DEBUG: Mostrar dados recebidos
    print(f"🔍 DEBUG preencher_documento - dados recebidos:")
    print(f"  - comarca_cliente: {dados.get('comarca_cliente')}")
    print(f"  - comarca_administradora: {dados.get('comarca_administradora')}")
    print(f"  - comarca_escolhida: {dados.get('comarca_escolhida')}")
    print(f"  - cpf: {dados.get('cpf')}")
    print(f"  - tipo_documento: {dados.get('tipo_documento')}")
    print(f"  - Total de chaves: {len(dados.keys())}")

    # substituição segura run-a-run, ignorando runs com desenhos/imagens
    def _repl(txt: str) -> str:
        if not txt:
            return txt
        out = txt
        for chave, valor in (dados or {}).items():
            v = "" if valor is None else str(valor)
            out = out.replace(f"{{{{ {chave} }}}}", v).replace(f"{{{{{chave}}}}}", v)
        return out

    def _run_tem_desenho(run) -> bool:
        el = run._element  # lxml element
        # namespaces-agnostic
        if el.xpath(".//*[local-name()='drawing']"):
            return True
        if el.xpath(".//*[local-name()='pict']"):
            return True
        if el.xpath(".//*[local-name()='shape']"):
            return True
        return False

    # parágrafos fora de tabelas
    for par in doc.paragraphs:
        for run in par.runs:
            if not _run_tem_desenho(run):
                run.text = _repl(run.text)

    # parágrafos dentro de tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for cel in linha.cells:
                for par in cel.paragraphs:
                    for run in par.runs:
                        if not _run_tem_desenho(run):
                            run.text = _repl(run.text)

    timestamp = now_sp().strftime('%Y%m%d%H%M%S')  # fuso America/Sao_Paulo
    slug = slugify(dados.get('nome') or dados.get('nome_cliente') or "cliente")
    nome_base = f"{nome_saida}_{slug}_{timestamp}"

    # ✅ CAMINHO ABSOLUTO CONFIÁVEL
    from app.utils.paths import get_documentos_dir
    pasta_saida = get_documentos_dir()

    caminho_docx = os.path.join(pasta_saida, f"{nome_base}.docx")
    caminho_pdf = os.path.join(pasta_saida, f"{nome_base}.pdf")

    doc.save(caminho_docx)
    converter_docx_para_pdf(caminho_docx, caminho_pdf, rasterizar_se_sem_img=False)

    if not os.path.exists(caminho_pdf):
        raise RuntimeError(f"PDF não gerado: {caminho_pdf}")

    return caminho_pdf

# ========================= HELPERS MODELOS =========================
def resolve_base_modelos(usuario_advogado: str) -> str:
    # Corrige caminho: modelos está em backend/modelos, não backend/app/modelos
    arquivo_atual = os.path.abspath(__file__)
    routes_dir = os.path.dirname(arquivo_atual)  # /backend/app/routes
    app_dir = os.path.dirname(routes_dir)  # /backend/app
    backend_dir = os.path.dirname(app_dir)  # /backend
    raiz = os.path.join(backend_dir, "modelos")  # /backend/modelos
    
    alvo = (usuario_advogado or "").strip().lower()
    if not alvo:
        return os.path.join(raiz, "vitor")
    pasta_adv = os.path.join(raiz, alvo)
    if os.path.isdir(pasta_adv):
        return pasta_adv
    return os.path.join(raiz, "vitor")

def localizar_modelo(pasta_base: str, tipo: str) -> str:
    tipo = (tipo or "").strip().lower()
    padrao = "contrato" if "contrato" in tipo else "procuracao"
    candidatos = glob.glob(os.path.join(pasta_base, "*.docx"))
    candidatos = [c for c in candidatos if padrao in os.path.basename(c).lower()]
    if candidatos:
        candidatos.sort()
        return candidatos[0]
    raise FileNotFoundError(f"Modelo '{padrao}' não encontrado em: {pasta_base}")

# ========================= FUNÇÕES EXPOSTAS PARA TESTES =========================
def gerar_documento_preview(dados: dict):
    """Função simples para testes que gera documentos sem salvar no banco"""
    try:
        usuario_advogado = (dados.get("usuario_advogado") or "").strip().lower()
        base_modelos = resolve_base_modelos(usuario_advogado)
        
        contrato_modelo = localizar_modelo(base_modelos, "contrato")
        procuracao_modelo = localizar_modelo(base_modelos, "procuracao")
        
        contrato_pdf_path = preencher_documento(contrato_modelo, dados, "contrato")
        procuracao_pdf_path = preencher_documento(procuracao_modelo, dados, "procuracao")
        
        return {
            "status": "ok",
            "contrato_pdf": os.path.basename(contrato_pdf_path),
            "procuracao_pdf": os.path.basename(procuracao_pdf_path),
            "nome_cliente": dados.get("nome") or dados.get("nome_cliente") or "",
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}

# ========================= ROTAS =========================
@router.post("/gerar-documentos-preview")
async def gerar_documentos_preview(dados: dict):
    try:
        usuario_advogado = (dados.get("usuario_advogado") or "").strip().lower()
        base_modelos = resolve_base_modelos(usuario_advogado)
        print(f"[Docs][Preview] Extrato {dados.get('extrato_id')} / advogado {usuario_advogado}")
        
        # 🔍 DEBUG: Ver campos essenciais ANTES
        print(f"🔍 [PREVIEW ANTES] cpf_cnpj={dados.get('cpf_cnpj')} | tipo_documento={dados.get('tipo_documento')}")
        
        # ✅ Detectar tipo_documento automaticamente se não vier
        if not dados.get('tipo_documento'):
            cpf_cnpj = dados.get('cpf_cnpj') or dados.get('cpf') or ''
            cpf_limpo = str(cpf_cnpj).replace('.', '').replace('-', '').replace('/', '').strip()
            if len(cpf_limpo) == 14:
                dados['tipo_documento'] = 'CNPJ'
                print(f"✅ [PREVIEW] tipo_documento detectado: CNPJ (14 dígitos)")
            else:
                dados['tipo_documento'] = 'CPF'
                print(f"✅ [PREVIEW] tipo_documento detectado: CPF (padrão)")
        
        print(f"🔍 [PREVIEW DEPOIS] cpf_cnpj={dados.get('cpf_cnpj')} | tipo_documento={dados.get('tipo_documento')}")

        contrato_modelo = localizar_modelo(base_modelos, "contrato")
        procuracao_modelo = localizar_modelo(base_modelos, "procuracao")

        contrato_pdf_path = preencher_documento(contrato_modelo, dados, "contrato")
        print(f"[Docs][Preview] Contrato preenchido -> {contrato_pdf_path}")

        procuracao_pdf_path = preencher_documento(procuracao_modelo, dados, "procuracao")
        print(f"[Docs][Preview] Procuração preenchida -> {procuracao_pdf_path}")

        return {
            "status": "ok",
            "contrato_pdf": os.path.basename(contrato_pdf_path),
            "procuracao_pdf": os.path.basename(procuracao_pdf_path),
            "nome_cliente": dados.get("nome") or dados.get("nome_cliente") or "",
            "telefone_cliente": dados.get("telefone") or dados.get("telefone_cliente") or "",
            "pasta_modelos": base_modelos,
        }
    except Exception as e:
        print(f"[preview] erro ao gerar documentos: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Falha ao gerar documentos (preview): {e}")

@router.post("/gerar-documentos")
def gerar_documentos(dados: dict, db: Session = Depends(get_db)):
    try:
        usuario_advogado = (dados.get("usuario_advogado") or "").strip().lower()
        if not usuario_advogado:
            raise HTTPException(status_code=400, detail="Usuário do advogado não informado")

        advogado = db.query(Advogado).filter(Advogado.usuario == usuario_advogado).first()
        if not advogado or not advogado.api_key_zapsign:
            raise HTTPException(status_code=404, detail="Advogado não encontrado ou sem chave ZapSign")

        base_modelos = resolve_base_modelos(usuario_advogado)
        print(f"[Docs][Geração] extrato={dados.get('extrato_id')} advogado={usuario_advogado}")
        
        # 🔍 DEBUG COMPLETO: Ver TODOS os dados recebidos
        print(f"🔍 [DEBUG GERAR-DOCUMENTOS] Dados completos recebidos:")
        print(f"   - nome: {dados.get('nome')}")
        print(f"   - cpf: {dados.get('cpf')}")
        print(f"   - cpf_cnpj: {dados.get('cpf_cnpj')}")
        print(f"   - tipo_documento: {dados.get('tipo_documento')}")
        print(f"   - Total campos: {len(dados)}")
        
        contrato_modelo = localizar_modelo(base_modelos, "contrato")
        procuracao_modelo = localizar_modelo(base_modelos, "procuracao")

        contrato_pdf_path = preencher_documento(contrato_modelo, dados, "contrato")
        print(f"[Docs][Geração] Contrato preenchido -> {contrato_pdf_path}")
        procuracao_pdf_path = preencher_documento(procuracao_modelo, dados, "procuracao")
        print(f"[Docs][Geração] Procuração preenchida -> {procuracao_pdf_path}")

        contrato_pdf_nome = os.path.basename(contrato_pdf_path)
        procuracao_pdf_nome = os.path.basename(procuracao_pdf_path)
        extrato_id = dados.get("extrato_id")

        print("[Docs][Geração] Upload Drive...")
        contrato_url = upload_pdf_to_drive(contrato_pdf_path, contrato_pdf_nome)
        procuracao_url = upload_pdf_to_drive(procuracao_pdf_path, procuracao_pdf_nome)

        # Persiste cópias físicas em storage/Assinaturas/<extrato_id>
        if extrato_id:
            _persist_pdf_in_storage(str(extrato_id), contrato_pdf_path, label="contrato_pdf")
            _persist_pdf_in_storage(str(extrato_id), procuracao_pdf_path, label="procuracao_pdf")

        print("[Docs][Geração] Chamando ZapSign...")
        try:
            webhook_url = build_advogado_webhook_url(advogado)
        except RuntimeError as e:
            raise HTTPException(status_code=500, detail=f"Webhook ZapSign indisponível: {e}")
        link_assinatura = enviar_documentos_consolidados_para_assinatura(
            nome_cliente=dados.get("nome") or dados.get("nome_cliente") or "",
            telefone_cliente=dados.get("telefone") or dados.get("telefone_cliente") or "",
            caminho_contrato=contrato_pdf_path,
            caminho_procuracao=procuracao_pdf_path,
            api_key=advogado.api_key_zapsign,
            webhook_url=webhook_url,
            sandbox=False,  # Produção
        )

        if extrato_id:
            sess = SessionLocal()
            try:
                extrato = sess.get(Extrato, int(extrato_id))
                if extrato:
                    # 🔍 DEBUG: Log antes da modificação
                    print(f"[DEBUG] ANTES - Extrato {extrato_id}: advogado_nome='{extrato.advogado_nome}', advogado_oab='{extrato.advogado_oab}'")
                    
                    extrato.status_documento = "enviado"
                    # Usar UTC para SQLite - garantia de timezone correto
                    from app.core.time import now_utc_for_sqlite
                    extrato.enviado_em = extrato.enviado_em or now_utc_for_sqlite()
                    
                    # 🔧 CORREÇÃO: NUNCA sobrescrever advogado em extrato existente
                    # Os dados do advogado devem ser preservados como foram cadastrados originalmente
                    print(f"[DEBUG] PRESERVANDO ADVOGADO - não alterando dados de {extrato.advogado_nome}")
                    # (comentado para preservar advogado original)
                    # extrato.advogado_id = getattr(advogado, "id", None)
                    # extrato.advogado_nome = getattr(advogado, "nome", None)
                    # extrato.advogado_oab = getattr(advogado, "oab", None)
                    # extrato.advogado_email = getattr(advogado, "email", None)
                    # extrato.advogado_telefone = getattr(advogado, "telefone", None)
                    extrato.contrato_url = contrato_url
                    extrato.procuracao_url = procuracao_url
                    links_dict = {}
                    if isinstance(link_assinatura, dict):
                        links_dict = (link_assinatura.get("links") or {})
                        if not links_dict and link_assinatura.get("link_assinatura"):
                            links_dict = {"link_assinatura": link_assinatura["link_assinatura"]}
                        extrato.zapsign_bundle_id = link_assinatura.get("bundle_id")
                        extrato.zapsign_contrato_id = link_assinatura.get("contrato_id")
                        extrato.zapsign_procuracao_id = link_assinatura.get("procuracao_id")
                    else:
                        links_dict = {"link_assinatura": str(link_assinatura)}
                    atual = dict(extrato.zapsign_links or {})
                    atual.update(links_dict or {})
                    extrato.zapsign_links = atual
                    sess.commit()
            except Exception as e:
                sess.rollback()
                print(f"[persistência] falha ao atualizar Extrato {extrato_id}: {e}")
            finally:
                sess.close()

        return {
            "mensagem": "Documentos gerados e enviados para assinatura com sucesso!",
            "contrato_pdf": contrato_pdf_nome,
            "procuracao_pdf": procuracao_pdf_nome,
            "contrato_url": contrato_url,
            "procuracao_url": procuracao_url,
            "nome_cliente": dados.get("nome", "") or dados.get("nome_cliente", ""),
            "telefone_cliente": dados.get("telefone", "") or dados.get("telefone_cliente", ""),
            "usuario_advogado": usuario_advogado,
            "link_assinatura": link_assinatura if isinstance(link_assinatura, str) else (link_assinatura.get("link_assinatura") or (link_assinatura.get("links") or {}).get("link_assinatura")),
            "pasta_modelos": base_modelos,
        }
    except HTTPException as e:
        raise e
    except Exception as e:
        print("❌ Erro ao gerar/enviar documentos:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/enviar-para-assinatura")
def enviar_para_assinatura(dados: dict, db: Session = Depends(get_db)):
    try:
        usuario_advogado = (dados.get("usuario_advogado") or "").strip().lower()
        if not usuario_advogado:
            raise HTTPException(status_code=400, detail="Usuário do advogado não informado")

        advogado = db.query(Advogado).filter(Advogado.usuario == usuario_advogado).first()
        if not advogado or not advogado.api_key_zapsign:
            raise HTTPException(status_code=404, detail="Advogado não encontrado ou sem chave ZapSign")

        contrato_pdf = dados.get("contrato_pdf")
        procuracao_pdf = dados.get("procuracao_pdf")
        if not contrato_pdf or not procuracao_pdf:
            raise HTTPException(status_code=400, detail="Documentos PDF não fornecidos")

        # ✅ CAMINHO ABSOLUTO CONFIÁVEL  
        from app.utils.paths import get_documentos_dir
        pasta_documentos = get_documentos_dir()
        caminho_contrato = os.path.join(pasta_documentos, contrato_pdf)
        caminho_procuracao = os.path.join(pasta_documentos, procuracao_pdf)
        print(f"[Docs][Envio] contrato={caminho_contrato} procuracao={caminho_procuracao} advogado={usuario_advogado}")
        if not os.path.exists(caminho_contrato) or not os.path.exists(caminho_procuracao):
            raise HTTPException(status_code=404, detail="Um ou mais arquivos não foram encontrados")

        extrato_id = dados.get("extrato_id")

        # Persiste em storage antes de enviar (compat com fluxo antigo)
        if extrato_id:
            _persist_pdf_in_storage(str(extrato_id), caminho_contrato, label="contrato_pdf")
            _persist_pdf_in_storage(str(extrato_id), caminho_procuracao, label="procuracao_pdf")

        print("[Docs][Envio] Upload Drive...")
        url_contrato = upload_pdf_to_drive(caminho_contrato, contrato_pdf)
        url_procuracao = upload_pdf_to_drive(caminho_procuracao, procuracao_pdf)

        print("[Docs][Envio] Chamando ZapSign...")
        link_assinatura = enviar_documentos_consolidados_para_assinatura(
            nome_cliente=dados.get("nome") or dados.get("nome_cliente") or "",
            telefone_cliente=dados.get("telefone") or dados.get("telefone_cliente") or "",
            caminho_contrato=caminho_contrato,
            caminho_procuracao=caminho_procuracao,
            api_key=advogado.api_key_zapsign,
            sandbox=False  # Produção
        )

        if extrato_id:
            sess = SessionLocal()
            try:
                extrato = sess.get(Extrato, int(extrato_id))
                if extrato:
                    extrato.status_documento = "enviado"
                    # Usar UTC para SQLite - garantia de timezone correto
                    from app.core.time import now_utc_for_sqlite
                    extrato.enviado_em = extrato.enviado_em or now_utc_for_sqlite()
                    extrato.contrato_url = url_contrato
                    extrato.procuracao_url = url_procuracao
                    links_dict = {"link_assinatura": link_assinatura} if isinstance(link_assinatura, str) else (link_assinatura.get("links") or {})
                    atual = dict(extrato.zapsign_links or {})
                    atual.update(links_dict or {})
                    extrato.zapsign_links = atual
                    sess.commit()
            except Exception as e:
                sess.rollback()
                print(f"[persistência] falha ao atualizar Extrato {extrato_id}: {e}")
            finally:
                sess.close()

        return {"mensagem": "Documentos enviados para assinatura com sucesso!", "link_assinatura": link_assinatura}
    except Exception as e:
        print("❌ Erro ao enviar documentos para assinatura:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
