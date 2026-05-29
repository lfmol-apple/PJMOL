from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Request, status, Header
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import text
from docx import Document
from datetime import datetime
import os
import secrets
import unicodedata
import re
import subprocess
import shutil
from typing import Optional, Tuple, Dict, Any, List
import json

from pydantic import BaseModel
from database import get_db
from app.models.advogado import Advogado
from app.models.extrato import Extrato
from app.utils.security import gerar_hash_senha
from app.utils.google_drive import upload_pdf_to_drive
from app.utils.zapsign import enviar_documentos_consolidados_para_assinatura, build_advogado_webhook_url
from app.core.time import now_sp


router = APIRouter(tags=["Advogados"])

# ✅ CAMINHOS ABSOLUTOS CONFIÁVEIS
from app.utils.paths import get_documentos_dir, get_modelos_dir, get_app_root, get_backend_root

DOCS_GERADOS_DIR = get_documentos_dir()
MODELOS_DIR = get_modelos_dir()
BASE_DIR = get_app_root()
BACKEND_ROOT = get_backend_root()

def slugify(texto: str) -> str:
    texto = unicodedata.normalize('NFD', texto)
    texto = texto.encode('ascii', 'ignore').decode('utf-8')
    texto = re.sub(r'\s+', '_', texto)
    texto = re.sub(r'[^a-zA-Z0-9_]', '', texto)
    return texto.lower()

def only_digits(s: Optional[str]) -> str:
    if not s:
        return ""
    return "".join(ch for ch in str(s) if ch.isdigit())

def _find_soffice_binary() -> Optional[str]:
    """
    Localiza o executável do LibreOffice considerando variáveis de ambiente e caminhos comuns em Linux/WSL.
    """
    candidates = [
        os.getenv("SOFFICE_PATH"),
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/usr/lib/libreoffice/program/soffice",
        "/usr/lib64/libreoffice/program/soffice",
        "/usr/local/bin/soffice",
    ]

    for cand in candidates:
        if not cand:
            continue
        if not os.path.isabs(cand):
            resolved = shutil.which(cand)
            if resolved:
                return resolved
            continue
        if os.path.isfile(cand) and os.access(cand, os.X_OK):
            return cand
    return None


def converter_docx_para_pdf(caminho_docx: str, pasta_saida: str):
    soffice_path = _find_soffice_binary()
    if not soffice_path:
        raise RuntimeError("LibreOffice (soffice) não encontrado no servidor.")

    os.makedirs(pasta_saida, exist_ok=True)
    cmd = [
        soffice_path,
        "--headless",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        pasta_saida,
        caminho_docx,
    ]
    try:
        result = subprocess.run(cmd, check=False, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except Exception as exc:
        raise RuntimeError(f"Erro ao executar LibreOffice: {exc}")

    if result.returncode != 0:
        stderr = result.stderr.decode("utf-8", errors="ignore") if result.stderr else ""
        raise RuntimeError(f"Erro na conversão com LibreOffice (retorno {result.returncode}): {stderr}")

def _substituir_placeholders_no_paragrafo(par, mapa: dict):
    for run in par.runs:
        txt = run.text or ""
        for chave, valor in mapa.items():
            token1 = f"{{{{ {chave} }}}}"
            token2 = f"{{{{{chave}}}}}"
            if token1 in txt or token2 in txt:
                txt = txt.replace(token1, valor).replace(token2, valor)
        run.text = txt

def _substituir_placeholders_no_documento(doc: Document, dados: dict):
    mapa = {k: ("" if v is None else str(v)) for k, v in dados.items()}
    for par in doc.paragraphs:
        _substituir_placeholders_no_paragrafo(par, mapa)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for cel in linha.cells:
                for par in cel.paragraphs:
                    _substituir_placeholders_no_paragrafo(par, mapa)

def preencher_documento(modelo_path: str, dados: dict, nome_saida: str) -> str:
    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Modelo não encontrado: {modelo_path}")

    doc = Document(modelo_path)
    _substituir_placeholders_no_documento(doc, dados)

    timestamp = now_sp().strftime('%Y%m%d%H%M%S')
    nome_cliente_slug = slugify(dados.get('nome', 'cliente'))
    nome_base = f"{nome_saida}_{nome_cliente_slug}_{timestamp}"

    caminho_docx = os.path.join(DOCS_GERADOS_DIR, f"{nome_base}.docx")
    caminho_pdf = os.path.join(DOCS_GERADOS_DIR, f"{nome_base}.pdf")

    doc.save(caminho_docx)
    converter_docx_para_pdf(caminho_docx, DOCS_GERADOS_DIR)

    if not os.path.exists(caminho_pdf):
        raise RuntimeError(f"❌ PDF não foi gerado: {caminho_pdf}")

    return os.path.basename(caminho_pdf)

def _resolver_modelos_do_advogado(adv: Advogado) -> Tuple[str, str]:
    usuario = (adv.usuario or "").strip().lower()
    if not usuario:
        raise HTTPException(status_code=400, detail="Advogado sem 'usuario' definido")

    pasta_usuario = os.path.join(MODELOS_DIR, usuario)
    candidatos = [
        (
            os.path.join(pasta_usuario, f"modelo_contrato_{usuario}.docx"),
            os.path.join(pasta_usuario, f"modelo_procuracao_{usuario}.docx"),
        ),
        (
            os.path.join(MODELOS_DIR, f"modelo_contrato_{usuario}.docx"),
            os.path.join(MODELOS_DIR, f"modelo_procuracao_{usuario}.docx"),
        ),
    ]

    for contrato_path, procuracao_path in candidatos:
        if os.path.exists(contrato_path) and os.path.exists(procuracao_path):
            return contrato_path, procuracao_path

    raise FileNotFoundError(
        f"Modelos não encontrados para '{usuario}'. "
        f"Esperado: modelo_contrato_{usuario}.docx e modelo_procuracao_{usuario}.docx "
        f"em {pasta_usuario} ou {MODELOS_DIR}"
    )

# ---------------------------
# Helpers (extras + audit)
# ---------------------------
def _coerce_extras(v) -> dict:
    if isinstance(v, dict):
        return dict(v)
    if isinstance(v, str):
        try:
            data = json.loads(v)
            return data if isinstance(data, dict) else {}
        except Exception:
            return {}
    return {}

def _audit_write(db: Session, *, action: str, message: str, extrato_id: int, actor_id: Optional[int], payload: Dict[str, Any]):
    try:
        q = text("""
            INSERT INTO audit_logs (action, message, extrato_id, actor_id, payload)
            VALUES (:action, :message, :extrato_id, :actor_id, :payload)
        """)
        db.execute(q, {
            "action": action,
            "message": message,
            "extrato_id": extrato_id,
            "actor_id": actor_id,
            "payload": json.dumps(payload, ensure_ascii=False)
        })
        db.commit()
    except Exception:
        try:
            db.rollback()
        except Exception:
            pass

# ---------------------------
# Gerar/servir documentos
# ---------------------------
@router.post("/gerar-documentos")
def gerar_documentos(dados: dict, db: Session = Depends(get_db)):
    try:
        usuario = dados.get("usuario_advogado")
        if not usuario:
            raise HTTPException(status_code=400, detail="Usuário do advogado não informado")

        advogado = db.query(Advogado).filter(Advogado.usuario == usuario).first()
        if not advogado or not advogado.api_key_zapsign:
            raise HTTPException(status_code=404, detail="Advogado não encontrado ou sem chave ZapSign")

        contrato_path, procuracao_path = _resolver_modelos_do_advogado(advogado)

        contrato_pdf = preencher_documento(contrato_path, dados, "contrato")
        procuracao_pdf = preencher_documento(procuracao_path, dados, "procuracao")

        caminho_contrato = os.path.join(DOCS_GERADOS_DIR, contrato_pdf)
        caminho_procuracao = os.path.join(DOCS_GERADOS_DIR, procuracao_pdf)

        url_contrato = upload_pdf_to_drive(caminho_contrato, os.path.basename(caminho_contrato))
        url_procuracao = upload_pdf_to_drive(caminho_procuracao, os.path.basename(caminho_procuracao))

        try:
            webhook_url = build_advogado_webhook_url(advogado)
        except RuntimeError as e:
            raise HTTPException(status_code=500, detail=f"Webhook ZapSign indisponível: {e}")

        link_assinatura = enviar_documentos_consolidados_para_assinatura(
            nome_cliente=dados.get("nome", ""),
            telefone_cliente=dados.get("telefone", ""),
            caminho_contrato=caminho_contrato,
            caminho_procuracao=caminho_procuracao,
            api_key=advogado.api_key_zapsign,
            webhook_url=webhook_url,
            sandbox=False,  # Produção
        )

        return {
            "mensagem": "Documentos gerados e enviados com sucesso!",
            "contrato_pdf": contrato_pdf,
            "procuracao_pdf": procuracao_pdf,
            "url_contrato": url_contrato,
            "url_procuracao": url_procuracao,
            "link_assinatura": link_assinatura
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/documentos/{nome_arquivo}")
def servir_documento(nome_arquivo: str):
    caminho = os.path.join(DOCS_GERADOS_DIR, nome_arquivo)
    if not os.path.exists(caminho):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")

    if nome_arquivo.endswith(".pdf"):
        media_type = "application/pdf"
    elif nome_arquivo.endswith(".docx"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        media_type = "application/octet-stream"

    return FileResponse(caminho, media_type=media_type)

# ---------------------------
# CRUD Advogados
# ---------------------------
class AtualizarChaveZapSign(BaseModel):
    api_key_zapsign: str

@router.get("/advogados/usuario/{usuario}")
def consultar_advogado_por_usuario(usuario: str, db: Session = Depends(get_db)):
    advogado = db.query(Advogado).filter(Advogado.usuario == usuario).first()
    if not advogado:
        raise HTTPException(status_code=404, detail="Advogado não encontrado")
    return {
        "id": advogado.id,
        "nome_completo": advogado.nome_completo,
        "usuario": advogado.usuario,
        "oab": advogado.oab,
        "email": advogado.email,
        "telefone": advogado.telefone,
        "api_key_zapsign": advogado.api_key_zapsign
    }

@router.post("/advogados/", status_code=status.HTTP_201_CREATED)
def criar_advogado(
    nome_completo: str = Form(...),
    oab: str = Form(...),
    email: str = Form(...),
    telefone: str = Form(""),
    usuario: str = Form(...),
    senha: str = Form(...),
    api_key_zapsign: str = Form(...),
    contrato: UploadFile = File(None),
    procuracao: UploadFile = File(None),
    db: Session = Depends(get_db),
):
    # NORMALIZAÇÕES
    usuario_norm = usuario.strip().lower()
    oab_norm = oab.strip().upper()
    email_norm = email.strip().lower()

    # 1) usuário deve ser único
    if db.query(Advogado).filter(Advogado.usuario == usuario_norm).first():
        raise HTTPException(status_code=400, detail="Usuário já existe")

    # 2) OAB deve ser única (REGRA PRINCIPAL)
    if db.query(Advogado).filter(Advogado.oab == oab_norm).first():
        raise HTTPException(status_code=409, detail="Já existe advogado com esta OAB")

    # (NÃO validamos e-mail como único — pode repetir)
    novo_advogado = Advogado(
        nome_completo=nome_completo.strip().upper(),
        oab=oab_norm,
        email=email_norm,
        telefone=only_digits(telefone),
        usuario=usuario_norm,
        senha_hash=gerar_hash_senha(senha),
        api_key_zapsign=api_key_zapsign.strip(),
    )
    db.add(novo_advogado)
    db.commit()
    db.refresh(novo_advogado)

    # salvar modelos, se enviados
    pasta_usuario = os.path.join(MODELOS_DIR, novo_advogado.usuario)
    os.makedirs(pasta_usuario, exist_ok=True)

    if contrato:
        caminho = os.path.join(pasta_usuario, f"modelo_contrato_{novo_advogado.usuario}.docx")
        with open(caminho, "wb") as f:
            f.write(contrato.file.read())

    if procuracao:
        caminho = os.path.join(pasta_usuario, f"modelo_procuracao_{novo_advogado.usuario}.docx")
        with open(caminho, "wb") as f:
            f.write(procuracao.file.read())

    return {"mensagem": "Advogado cadastrado com sucesso!", "id": novo_advogado.id}

@router.get("/advogados")
@router.get("/advogados/")
def listar_advogados(db: Session = Depends(get_db)):
    # Lista de advogados inativos (saíram da equipe)
    ADVOGADOS_INATIVOS = ['melisa']
    
    advogados = db.query(Advogado).all()
    return [
        {
            "id": adv.id,
            "nome_completo": adv.nome_completo,
            "usuario": adv.usuario,
            "email": adv.email,
            "telefone": adv.telefone,
            "oab": adv.oab,
            "api_key_zapsign": adv.api_key_zapsign,
            "ativo": adv.usuario not in ADVOGADOS_INATIVOS
        }
        for adv in advogados
    ]


# ========= Criar advogado a partir do template padrão do escritório =========
class CriarAdvogadoComTemplateIn(BaseModel):
    nome_completo: str
    oab: str
    email: str
    telefone: str = ""
    usuario: str
    senha: str
    api_key_zapsign: str
    webhook_path_token: Optional[str] = None
    genero: str = "M"   # "M" masculino / "F" feminino




PADRAO_CONTRATO   = os.path.join(MODELOS_DIR, '_padrao', 'modelo_contrato_padrao.docx')
PADRAO_PROCURACAO = os.path.join(MODELOS_DIR, '_padrao', 'modelo_procuracao_padrao.docx')


def _gerar_tokens(adv, genero: str) -> dict:
    """Gera o mapa de tokens → valores para preencher o template padrão."""
    feminino = genero.strip().upper() == "F"

    nome = (adv.nome_completo or "").strip().upper()

    digits = re.sub(r"\D", "", adv.oab or "")
    letters = "".join(c for c in (adv.oab or "") if c.isalpha())[:2].upper()
    oab_proc     = f"OAB/{letters} {digits[:-3]}.{digits[-3:]}" if len(digits) > 3 else f"OAB/{letters} {digits}"
    oab_contrato = f"OAB/{letters} {digits}"

    d = re.sub(r"\D", "", adv.telefone or "")
    if len(d) == 11:
        tel = f"({d[:2]}) {d[2:7]}-{d[7:]}"
    elif len(d) == 10:
        tel = f"({d[:2]}) {d[2:6]}-{d[6:]}"
    else:
        tel = d

    return {
        "__ADVOGADO_NOME__":    nome,
        "__ADV_OAB_CONTRATO__": oab_contrato,
        "__ADV_OAB_PROC__":     oab_proc,
        "__ADVOGADO_TELEFONE__": tel,
        "__DR_DRA__":           "Dra" if feminino else "Dr",
        "__O_A_ART__":          "a"   if feminino else "o",
        "__CONTRATADO_A__":     "CONTRATADA" if feminino else "CONTRATADO",
        "__DENOMINADO_A__":     "denominada" if feminino else "denominado",
        "__INSCRITO_A__":       "inscrita"   if feminino else "inscrito",
        "__PROCURADOR_A__":     "procuradora" if feminino else "procurador",
    }


def _aplicar_tokens_no_paragrafo(par, tokens: dict):
    for run in par.runs:
        for tok, val in tokens.items():
            if tok in (run.text or ""):
                run.text = run.text.replace(tok, val)
    full = par.text
    if any(tok in full for tok in tokens):
        novo = full
        for tok, val in tokens.items():
            novo = novo.replace(tok, val)
        if par.runs:
            par.runs[0].text = novo
            for run in par.runs[1:]:
                run.text = ""


def _aplicar_tokens_template(caminho_docx: str, tokens: dict):
    doc = Document(caminho_docx)
    for par in doc.paragraphs:
        _aplicar_tokens_no_paragrafo(par, tokens)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    _aplicar_tokens_no_paragrafo(par, tokens)
    doc.save(caminho_docx)



@router.post("/advogados/com-template/", status_code=status.HTTP_201_CREATED)
def criar_advogado_com_template(
    dados: CriarAdvogadoComTemplateIn,
    db: Session = Depends(get_db),
):
    # Verificar templates padrão
    for tp in (PADRAO_CONTRATO, PADRAO_PROCURACAO):
        if not os.path.isfile(tp):
            raise HTTPException(status_code=500, detail=f"Template padrão não encontrado: {tp}")

    usuario_norm = dados.usuario.strip().lower()
    oab_norm = dados.oab.strip().upper()
    email_norm = dados.email.strip().lower()

    if db.query(Advogado).filter(Advogado.usuario == usuario_norm).first():
        raise HTTPException(status_code=400, detail="Usuário já existe")

    if db.query(Advogado).filter(Advogado.oab == oab_norm).first():
        raise HTTPException(status_code=409, detail="Já existe advogado com esta OAB")

    novo_advogado = Advogado(
        nome_completo=dados.nome_completo.strip().upper(),
        oab=oab_norm,
        email=email_norm,
        telefone=only_digits(dados.telefone),
        usuario=usuario_norm,
        senha_hash=gerar_hash_senha(dados.senha),
        api_key_zapsign=dados.api_key_zapsign.strip(),
        webhook_path_token=(dados.webhook_path_token or "").strip() or secrets.token_hex(16),
    )
    db.add(novo_advogado)
    db.commit()
    db.refresh(novo_advogado)

    try:
        pasta_novo = os.path.join(MODELOS_DIR, novo_advogado.usuario)
        os.makedirs(pasta_novo, exist_ok=True)

        caminho_novo_contrato = os.path.join(pasta_novo, f"modelo_contrato_{novo_advogado.usuario}.docx")
        caminho_nova_procuracao = os.path.join(pasta_novo, f"modelo_procuracao_{novo_advogado.usuario}.docx")

        shutil.copy2(PADRAO_CONTRATO, caminho_novo_contrato)
        shutil.copy2(PADRAO_PROCURACAO, caminho_nova_procuracao)

        tokens = _gerar_tokens(novo_advogado, dados.genero)
        _aplicar_tokens_template(caminho_novo_contrato, tokens)
        _aplicar_tokens_template(caminho_nova_procuracao, tokens)
    except Exception as exc:
        db.delete(novo_advogado)
        db.commit()
        raise HTTPException(status_code=500, detail=f"Falha ao gerar templates: {exc}")

    return {
        "mensagem": f"Advogado '{novo_advogado.nome_completo}' cadastrado com sucesso.",
        "id": novo_advogado.id,
        "nome_completo": novo_advogado.nome_completo,
        "usuario": novo_advogado.usuario,
        "oab": novo_advogado.oab,
        "email": novo_advogado.email,
        "telefone": novo_advogado.telefone,
        "api_key_zapsign": novo_advogado.api_key_zapsign,
    }


@router.get("/advogados/{advogado_id}")
def get_advogado_por_id(advogado_id: int, db: Session = Depends(get_db)):
    advogado = db.query(Advogado).filter(Advogado.id == advogado_id).first()
    if not advogado:
        raise HTTPException(status_code=404, detail="Advogado não encontrado")
    return {
        "id": advogado.id,
        "nome_completo": advogado.nome_completo,
        "usuario": advogado.usuario,
        "email": advogado.email,
        "telefone": advogado.telefone,
        "oab": advogado.oab,
        "api_key_zapsign": advogado.api_key_zapsign
    }

@router.put("/advogados/usuario/{usuario}")
def atualizar_api_key(usuario: str, dados: AtualizarChaveZapSign, db: Session = Depends(get_db)):
    advogado = db.query(Advogado).filter(Advogado.usuario == usuario).first()
    if not advogado:
        raise HTTPException(status_code=404, detail="Advogado não encontrado")
    advogado.api_key_zapsign = (dados.api_key_zapsign or "").strip()
    db.commit()
    db.refresh(advogado)
    return {
        "mensagem": "Chave ZapSign atualizada com sucesso",
        "usuario": advogado.usuario,
        "nova_api_key": advogado.api_key_zapsign
    }

# ===== NOVOS =====
@router.put("/advogados/{advogado_id}")
async def atualizar_advogado(
    advogado_id: int,
    request: Request,
    # se vier multipart/form-data
    nome_completo: Optional[str] = Form(None),
    oab: Optional[str] = Form(None),
    email: Optional[str] = Form(None),
    telefone: Optional[str] = Form(None),
    usuario: Optional[str] = Form(None),
    api_key_zapsign: Optional[str] = Form(None),
    senha: Optional[str] = Form(None),
    db: Session = Depends(get_db),
):
    adv = db.query(Advogado).filter(Advogado.id == advogado_id).first()
    if not adv:
        raise HTTPException(status_code=404, detail="Advogado não encontrado")

    payload: Dict[str, Any] = {}
    ct = request.headers.get("content-type", "")
    if ct.startswith("application/json"):
        try:
            payload = await request.json()
        except Exception:
            payload = {}
    else:
        payload = {
            "nome_completo": nome_completo,
            "oab": oab,
            "email": email,
            "telefone": telefone,
            "usuario": usuario,
            "api_key_zapsign": api_key_zapsign,
            "senha": senha,
        }

    # Validação de usuário único (se alterar)
    if payload.get("usuario") is not None:
        novo_usuario = str(payload["usuario"]).strip().lower()
        if novo_usuario and novo_usuario != adv.usuario:
            if db.query(Advogado).filter(Advogado.usuario == novo_usuario).first():
                raise HTTPException(status_code=400, detail="Usuário já em uso por outro advogado")
            adv.usuario = novo_usuario

    # Validação de OAB única (se alterar) — REGRA PRINCIPAL
    if payload.get("oab") is not None:
        nova_oab = str(payload["oab"]).strip().upper()
        if nova_oab and nova_oab != adv.oab:
            conflito = db.query(Advogado).filter(Advogado.oab == nova_oab).first()
            if conflito:
                raise HTTPException(status_code=409, detail="Já existe advogado com esta OAB")
            adv.oab = nova_oab

    # E-mail pode repetir — apenas normalizamos
    if payload.get("email") is not None:
        adv.email = str(payload["email"]).strip().lower()

    if payload.get("nome_completo") is not None:
        adv.nome_completo = str(payload["nome_completo"]).strip().upper()
    if payload.get("telefone") is not None:
        adv.telefone = only_digits(payload["telefone"])
    if payload.get("api_key_zapsign") is not None:
        adv.api_key_zapsign = str(payload["api_key_zapsign"]).strip()
    if payload.get("senha"):
        adv.senha_hash = gerar_hash_senha(str(payload["senha"]))

    db.commit()
    db.refresh(adv)
    return {
        "id": adv.id,
        "nome_completo": adv.nome_completo,
        "usuario": adv.usuario,
        "email": adv.email,
        "telefone": adv.telefone,
        "oab": adv.oab,
        "api_key_zapsign": adv.api_key_zapsign,
        "mensagem": "Advogado atualizado com sucesso"
    }

@router.delete("/advogados/id/{advogado_id}", status_code=status.HTTP_200_OK)
def deletar_advogado_por_id(advogado_id: int, db: Session = Depends(get_db)):
    adv = db.query(Advogado).filter(Advogado.id == advogado_id).first()
    if not adv:
        raise HTTPException(status_code=404, detail="Advogado não encontrado")
    db.delete(adv)
    db.commit()
    return {"mensagem": f"Advogado com ID {advogado_id} foi excluído com sucesso."}

# ========= NOVO ENDPOINT: registrar número do processo (admin/interno) =========
class NumeroProcessoIn(BaseModel):
    numero_processo: str

@router.post("/advogados/extratos/{extrato_id}/numero-processo")
def set_numero_processo_interno(
    extrato_id: int,
    body: NumeroProcessoIn,
    db: Session = Depends(get_db),
    x_usuario_id: Optional[int] = Header(None, alias="X-Usuario-Id"),
):
    """
    Registra/atualiza o número do processo e carimba:
      - extras.numero_processo_set_at (se ainda não houver)
      - extras.numero_processo_last_set_at (sempre atualiza)
    Grava em audit_logs: action='process.number.set'
    """
    ex = db.query(Extrato).filter(Extrato.id == extrato_id).first()
    if not ex:
        raise HTTPException(status_code=404, detail="Extrato não encontrado")

    ex.numero_processo = (body.numero_processo or "").strip()
    if not ex.numero_processo:
        raise HTTPException(status_code=400, detail="numero_processo vazio")

    extras = _coerce_extras(ex.extras)
    now_iso = now_sp().isoformat()

    # preserva o primeiro carimbo; sempre atualiza o "last"
    if not extras.get("numero_processo_set_at"):
        extras["numero_processo_set_at"] = now_iso
    extras["numero_processo_last_set_at"] = now_iso

    ex.extras = extras
    
    # Atualiza timestamp do número do processo (finaliza timer do advogado)
    from app.routes.uploads_clean import _set_numero_processo_timestamp
    _set_numero_processo_timestamp(ex, db)
    
    db.add(ex)
    db.commit()
    db.refresh(ex)

    # Audit
    try:
        _audit_write(
            db,
            action="process.number.set",
            message="Número do processo definido/atualizado (interno).",
            extrato_id=extrato_id,
            actor_id=x_usuario_id if isinstance(x_usuario_id, int) else None,
            payload={"numero_processo": ex.numero_processo, "at": now_iso}
        )
    except Exception:
        pass

    return {
        "ok": True,
        "extrato_id": ex.id,
        "numero_processo": ex.numero_processo,
        "numero_processo_set_at": extras.get("numero_processo_set_at"),
        "numero_processo_last_set_at": extras.get("numero_processo_last_set_at"),
    }
